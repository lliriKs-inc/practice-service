from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT = (
    Path(__file__).resolve().parents[2]
    / "templates"
    / "documents"
    / "notice.docx"
)


def set_run_font(run, *, size=14, bold=False):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(document, text="", *, align=None, bold=False, before=0, after=6):
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.15
    set_run_font(paragraph.add_run(text), bold=bold)
    return paragraph


def build_notice():
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(14)

    add_paragraph(
        document,
        "ИЗВЕЩЕНИЕ",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        after=3,
    )
    add_paragraph(
        document,
        "о прохождении производственной практики",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        after=18,
    )

    paragraph = add_paragraph(document, after=10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_run_font(paragraph.add_run("Настоящим подтверждается, что студент(ка) "))
    set_run_font(paragraph.add_run("{{student_fio}}"), bold=True)
    set_run_font(paragraph.add_run(" группы "))
    set_run_font(paragraph.add_run("{{group}}"), bold=True)
    set_run_font(paragraph.add_run(" проходит производственную практику в организации ИП Езуб Антон Сергеевич."))

    paragraph = add_paragraph(document, after=10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_run_font(paragraph.add_run("Срок практики: с "))
    set_run_font(paragraph.add_run("{{practice_start}}"), bold=True)
    set_run_font(paragraph.add_run(" по "))
    set_run_font(paragraph.add_run("{{practice_end}}"), bold=True)
    set_run_font(paragraph.add_run("."))

    paragraph = add_paragraph(document, after=24)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_run_font(paragraph.add_run("Тема практики: "))
    set_run_font(paragraph.add_run("{{practice_topic}}"), bold=True)
    set_run_font(paragraph.add_run("."))

    add_paragraph(document, "Руководитель практики от организации", after=18)
    add_paragraph(document, "__________________ / Езуб А. С. /", after=18)
    add_paragraph(document, "М. П. (при наличии)", after=18)
    add_paragraph(document, "Дата: «___» __________ 20___ г.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)


if __name__ == "__main__":
    build_notice()
